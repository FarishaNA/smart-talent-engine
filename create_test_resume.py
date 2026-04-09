
from docx import Document
import os

def create_test_resume(filename):
    doc = Document()
    doc.add_heading('Farisha N A', 0)
    
    doc.add_heading('Education', level=1)
    doc.add_paragraph('B.Tech in Computer Science, XYZ University, 2022-2025')
    
    doc.add_heading('Experience', level=1)
    doc.add_paragraph('Summer Intern, Tech Corp, June 2024 - August 2024')
    
    doc.add_heading('Projects', level=1)
    doc.add_paragraph('Smart Talent Engine, 2024 - Present')
    doc.add_paragraph('Building a resume screening system using Python and Django.')
    
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Python, Django, REST APIs, JavaScript')
    
    doc.save(filename)
    print(f"Created test resume: {filename}")

if __name__ == "__main__":
    create_test_resume("farisha_resume.docx")
