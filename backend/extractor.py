"""
extractor.py
============
Multi-format resume text extraction.

Supports: PDF, DOCX, DOC, JPG, PNG, JPEG, WEBP

Primary parser: Docling
  - Handles two-column layouts correctly by detecting column boundaries
  - Handles tables without merging unrelated cells
  - Handles sidebars as separate text blocks
  - Produces structured markdown preserving section hierarchy
  - Falls back to pdfplumber if Docling fails

Image parser: EasyOCR
  - Handles scanned PDFs and image resumes
  - Better than pytesseract on low-quality and rotated text
  - Falls back to pytesseract if EasyOCR fails

DOCX parser: python-docx
  - Direct XML extraction, preserves paragraph structure
  - Handles tables and lists correctly
"""

import re
import uuid
from pathlib import Path
from datetime import datetime
from typing import Optional


# ─── SECTION DETECTION ────────────────────────────────────────────────────────

SECTION_MAP = {
    "experience": [
        "work experience", "professional experience", "employment history",
        "experience", "career history", "work history", "employment"
    ],
    "education": [
        "education", "academic background", "educational qualification",
        "academic qualification", "qualifications", "academic history"
    ],
    "projects": [
        "projects", "personal projects", "academic projects",
        "side projects", "portfolio", "key projects", "notable projects",
        "technical projects", "project experience"
    ],
    "skills": [
        "skills", "technical skills", "core competencies", "technologies",
        "tech stack", "tools", "programming languages", "frameworks",
        "technical expertise", "competencies", "proficiencies"
    ],
    "certifications": [
        "certifications", "certificates", "courses", "training",
        "professional development", "achievements", "awards"
    ],
    "summary": [
        "summary", "objective", "profile", "about", "overview",
        "professional summary", "career objective", "about me"
    ]
}


def detect_sections(text: str) -> dict:
    """
    Split resume text into labelled sections.
    Handles both uppercase and mixed-case headers.
    Returns dict of {section_name: section_text}.
    """
    sections = {"general": []}
    current = "general"
    lines = text.split('\n')

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Check if this line is a section header
        # Headers are typically short (<60 chars) and match known patterns
        lower = stripped.lower()
        # Remove common markdown symbols from header detection
        clean_lower = re.sub(r'^#+\s*', '', lower).strip()

        matched_section = None
        if len(stripped) < 60:
            for section_name, keywords in SECTION_MAP.items():
                for keyword in keywords:
                    if clean_lower == keyword or clean_lower.startswith(keyword + ' '):
                        matched_section = section_name
                        break
                if matched_section:
                    break

        if matched_section:
            # Save current buffer
            if sections.get(current):
                pass  # already accumulated
            current = matched_section
            sections.setdefault(current, [])
        else:
            sections.setdefault(current, []).append(stripped)

    return {k: '\n'.join(v).strip() for k, v in sections.items() if v}


# ─── YEAR AND EXPERIENCE EXTRACTION ───────────────────────────────────────────

def extract_year_ranges(text: str) -> list[dict]:
    """
    Extract year ranges from text.

    CRITICAL FIX: Only extract year ranges from EXPERIENCE and PROJECTS
    sections, not from EDUCATION sections. This prevents education years
    (e.g. 2022-2025 college) from being counted as work experience.

    Patterns matched:
    - 2021 - 2024
    - 2021-2024
    - Jan 2021 - Dec 2024
    - 2021 to 2024
    - 2021 - Present
    - 2021 - Current
    """
    current_year = datetime.now().year

    # Pattern for year ranges
    pattern = (
        r'(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|'
        r'jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|'
        r'dec(?:ember)?)?\s*'
        r'(20\d{2})\s*'
        r'(?:–|-|to|—)\s*'
        r'(20\d{2}|present|current|now|ongoing)'
    )

    ranges = []
    for m in re.finditer(pattern, text.lower()):
        start = int(m.group(1))
        end_raw = m.group(2).strip()

        if end_raw in ('present', 'current', 'now', 'ongoing'):
            end = current_year
        else:
            try:
                end = int(end_raw)
            except ValueError:
                continue

        # Sanity checks
        if start < 1990 or start > current_year:
            continue
        if end < start:
            continue
        end = min(end, current_year)  # Cap future dates at current year
        duration = end - start

        if duration < 0 or duration > 15:
            continue  # Skip implausible ranges

        ranges.append({
            "start": start,
            "end": end,
            "duration": duration,
            "position": m.start()
        })

    return ranges


def calculate_experience(sections: dict) -> dict:
    """
    CRITICAL FIX: Calculate experience ONLY from experience and projects
    sections. Education section years must NOT be counted as experience.

    This fixes the bug where a student with education dates
    2022-2025 was getting 3 years of 'experience' counted.
    """
    # Only look for year ranges in these sections
    experience_text = (
        sections.get("experience", "") + "\n" +
        sections.get("projects", "") + "\n" +
        sections.get("general", "")
    )

    # Explicitly exclude education
    # DO NOT pass sections.get("education") to extract_year_ranges

    year_ranges = extract_year_ranges(experience_text)

    if not year_ranges:
        return {
            "year_ranges": [],
            "total_experience_years": 0,
            "most_recent_year": None,
            "has_professional_experience": False
        }

    total = sum(r["duration"] for r in year_ranges)
    most_recent = max(r["end"] for r in year_ranges)
    current_year = datetime.now().year

    return {
        "year_ranges": year_ranges,
        "total_experience_years": min(total, 30),  # Cap at 30
        "most_recent_year": most_recent,
        "has_professional_experience": total > 0,
        "is_recent": most_recent >= current_year - 1
    }


# ─── KEYWORD STUFFING DETECTION ───────────────────────────────────────────────

def detect_keyword_stuffing(sections: dict, raw_text: str) -> bool:
    """
    Detect resume keyword stuffing.
    Flag if: skills section is >35% of total words
    AND experience/projects section has <80 words total.
    Both conditions must be true — avoids false positives.
    """
    skills_text = sections.get("skills", "") + sections.get("technical skills", "")
    exp_text = sections.get("experience", "") + sections.get("projects", "")

    total_words = len(raw_text.split())
    skills_words = len(skills_text.split())
    exp_words = len(exp_text.split())

    if total_words == 0:
        return False

    high_skill_density = (skills_words / total_words) > 0.35
    thin_experience = exp_words < 80

    return high_skill_density and thin_experience


# ─── NAME EXTRACTION ──────────────────────────────────────────────────────────

def extract_name(text: str) -> str:
    """
    Extract candidate name from resume.

    Handles:
    - Standard: "John Smith"
    - With initials: "Farisha N A" (South Indian format)
    - Reversed: "Smith, John"
    - With titles: drops Dr., Mr., Ms. prefixes

    Strategy: Check first 8 non-empty lines for a name-like pattern.
    Skip lines containing contact info markers.
    """
    SKIP_MARKERS = [
        '@', 'http', 'linkedin', 'github', 'phone', 'email',
        '+91', '+1', 'tel:', 'mobile', '|', '•', 'objective',
        'summary', 'profile', 'cv', 'resume', 'curriculum'
    ]

    TITLE_PREFIXES = ['dr.', 'mr.', 'ms.', 'mrs.', 'prof.', 'er.']

    lines = [l.strip() for l in text.split('\n') if l.strip()]
    checked = 0

    for line in lines:
        if checked >= 10:
            break

        lower = line.lower()

        # Skip lines with contact/header markers
        if any(marker in lower for marker in SKIP_MARKERS):
            checked += 1
            continue

        # Skip very short or very long lines
        words = line.split()
        if len(words) < 2 or len(words) > 7:
            checked += 1
            continue

        # Remove title prefixes
        clean_words = [w for w in words
                      if w.lower().replace('.', '') + '.' not in TITLE_PREFIXES]
        if not clean_words:
            checked += 1
            continue

        # Check if mostly alphabetic (allows dots and spaces for initials)
        clean_line = ' '.join(clean_words)
        alpha_ratio = sum(
            c.isalpha() or c in ' .'
            for c in clean_line
        ) / max(len(clean_line), 1)

        if alpha_ratio > 0.85:
            return clean_line.strip()

        checked += 1

    return "Unknown"


# ─── CANDIDATE TYPE CLASSIFICATION ────────────────────────────────────────────

def classify_candidate_type(exp_data: dict, sections: dict) -> dict:
    """
    Classify candidate into experience tier based on EXPERIENCE data only.
    Education section dates are explicitly excluded.
    """
    current_year = datetime.now().year
    total_exp = exp_data.get("total_experience_years", 0)
    most_recent = exp_data.get("most_recent_year")
    has_exp = exp_data.get("has_professional_experience", False)

    # Gap detection: no professional activity in 2+ years
    if has_exp and most_recent and most_recent < current_year - 1:
        gap = current_year - most_recent
        return {
            "type": "career_gap",
            "label": f"Career Gap ({gap}yr)",
            "badge_class": "bg-orange-100 text-orange-800",
            "detail": f"Last professional activity: {most_recent}",
            "scoring_note": "Verify gap reason before filtering",
            "score_modifier": 0.90
        }

    # Fresher: no professional experience at all
    if not has_exp or total_exp == 0:
        project_text = (
            sections.get("projects", "") +
            sections.get("personal projects", "") +
            sections.get("academic projects", "")
        )
        strong_projects = len(project_text.split()) > 150

        return {
            "type": "fresher",
            "label": "Fresher",
            "badge_class": "bg-blue-100 text-blue-800",
            "detail": (
                "No professional experience — strong project portfolio"
                if strong_projects
                else "No professional experience — limited project evidence"
            ),
            "scoring_note": (
                "Score reflects academic context. "
                "Evaluate on project depth, not experience years."
            ),
            "score_modifier": 1.0
        }

    # Tiered by years
    if total_exp <= 2:
        return {
            "type": "junior",
            "label": f"Junior ({total_exp}yr exp)",
            "badge_class": "bg-green-100 text-green-800",
            "detail": f"{total_exp} years professional experience",
            "scoring_note": None,
            "score_modifier": 1.0
        }
    elif total_exp <= 5:
        return {
            "type": "mid",
            "label": f"Mid-level ({total_exp}yr exp)",
            "badge_class": "bg-teal-100 text-teal-800",
            "detail": f"{total_exp} years professional experience",
            "scoring_note": None,
            "score_modifier": 1.0
        }
    elif total_exp <= 10:
        return {
            "type": "senior",
            "label": f"Senior ({total_exp}yr exp)",
            "badge_class": "bg-purple-100 text-purple-800",
            "detail": f"{total_exp} years professional experience",
            "scoring_note": None,
            "score_modifier": 1.0
        }
    else:
        return {
            "type": "lead",
            "label": f"Lead/Principal ({total_exp}yr exp)",
            "badge_class": "bg-indigo-100 text-indigo-800",
            "detail": f"{total_exp}+ years professional experience",
            "scoring_note": "May be overqualified for junior roles",
            "score_modifier": 1.0
        }


# ─── SKILL VERIFICATION ───────────────────────────────────────────────────────

def verify_skills_against_evidence(sections: dict, graph: dict) -> dict:
    """
    Cross-reference skills section claims against experience and projects.

    Three categories:
    - verified: skill appears in BOTH skills section AND experience/projects
    - claimed_only: skill appears ONLY in skills section (possible padding)
    - demonstrated_not_claimed: skill in experience/projects but NOT skills section
      (hidden capability — candidate is underreporting themselves)

    verification_score: ratio of verified to total claimed skills
    """
    # Import here to avoid circular imports
    try:
        from alias_matcher import match_skills_to_graph
    except ImportError:
        return {"verified": [], "claimed_only": [], "demonstrated_not_claimed": [],
                "verification_score": 0.0, "summary": "Skill verification unavailable"}

    skills_text = (
        sections.get("skills", "") + " " +
        sections.get("technical skills", "")
    ).strip()

    evidence_text = (
        sections.get("experience", "") + " " +
        sections.get("professional experience", "") + " " +
        sections.get("projects", "") + " " +
        sections.get("personal projects", "") + " " +
        sections.get("academic projects", "")
    ).strip()

    claimed = {s["node_id"]: s for s in match_skills_to_graph(skills_text, graph)}
    demonstrated = {s["node_id"]: s for s in match_skills_to_graph(evidence_text, graph)}

    verified, claimed_only, hidden = [], [], []

    for nid, skill in claimed.items():
        if nid in demonstrated:
            verified.append({**skill, "verification": "verified"})
        else:
            claimed_only.append({**skill, "verification": "claimed_only"})

    for nid, skill in demonstrated.items():
        if nid not in claimed:
            hidden.append({**skill, "verification": "demonstrated_not_claimed"})

    total_claimed = len(claimed)
    verification_score = len(verified) / total_claimed if total_claimed > 0 else 0.0

    return {
        "verified": verified,
        "claimed_only": claimed_only,
        "demonstrated_not_claimed": hidden,
        "verification_score": round(verification_score, 2),
        "summary": (
            f"{len(verified)} verified, "
            f"{len(claimed_only)} claimed only, "
            f"{len(hidden)} hidden capabilities"
        )
    }


# ─── PRIMARY PARSERS ──────────────────────────────────────────────────────────

def parse_with_docling(file_path: str) -> tuple[str, float]:
    """
    Parse PDF using Docling.
    Docling handles two-column layouts, tables, and sidebars correctly.
    Returns (clean_text, confidence).
    """
    try:
        from docling.document_converter import DocumentConverter
        converter = DocumentConverter()
        result = converter.convert(file_path)
        markdown = result.document.export_to_markdown()

        # Clean markdown symbols while preserving structure
        clean = re.sub(r'^#{1,6}\s+', '\n', markdown, flags=re.MULTILINE)
        clean = re.sub(r'\*{1,2}([^*\n]+)\*{1,2}', r'\1', clean)
        clean = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', clean)
        clean = re.sub(r'\n{3,}', '\n\n', clean)
        clean = clean.strip()

        word_count = len(clean.split())
        confidence = min(1.0, word_count / 100)
        return clean, round(confidence, 2)

    except Exception as e:
        print(f"[Docling] Failed: {e}")
        return "", 0.0


def parse_with_pdfplumber(file_path: str) -> tuple[str, float]:
    """Fallback PDF parser using pdfplumber."""
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text(layout=True)
                if text:
                    parts.append(text)
        full = '\n'.join(parts)
        confidence = min(1.0, len(full.split()) / 100)
        return full, round(confidence, 2)
    except Exception as e:
        print(f"[pdfplumber] Failed: {e}")
        return "", 0.0


def parse_docx(file_path: str) -> tuple[str, float]:
    """Parse DOCX files using python-docx."""
    try:
        import docx as python_docx
        doc = python_docx.Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        full = '\n'.join(paragraphs)
        confidence = min(1.0, len(full.split()) / 100)
        return full, round(confidence, 2)
    except Exception as e:
        print(f"[python-docx] Failed: {e}")
        return "", 0.0


def parse_image(file_path: str) -> tuple[str, float]:
    """
    Parse image resumes (JPG, PNG, etc.) using EasyOCR with pytesseract fallback.
    """
    # Try EasyOCR first
    try:
        import easyocr
        reader = easyocr.Reader(['en'], gpu=False)
        result = reader.readtext(file_path)
        text = '\n'.join([item[1] for item in result])
        confidence = min(1.0, len(text.split()) / 50)
        if len(text.split()) > 20:
            return text, round(confidence, 2)
    except Exception as e:
        print(f"[EasyOCR] Failed: {e}")

    # Fallback to pytesseract
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img)
        confidence = min(1.0, len(text.split()) / 50)
        return text, round(confidence * 0.8, 2)  # Lower confidence for tesseract
    except Exception as e:
        print(f"[pytesseract] Failed: {e}")
        return "", 0.0


# ─── MAIN ENTRY POINT ─────────────────────────────────────────────────────────

def parse_resume(file_path: str) -> dict:
    """
    Parse a resume file and return structured data.

    Routing:
    - .pdf  → Docling (primary) → pdfplumber (fallback)
    - .docx/.doc → python-docx
    - .jpg/.jpeg/.png/.webp → EasyOCR → pytesseract
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    # Route to correct parser
    if ext == ".pdf":
        raw_text, confidence = parse_with_docling(file_path)
        if len(raw_text.split()) < 30:
            print("[Extractor] Docling produced little text, trying pdfplumber")
            raw_text, confidence = parse_with_pdfplumber(file_path)
        if len(raw_text.split()) < 10:
            # Might be a scanned PDF — try OCR
            print("[Extractor] PDF has no text layer, trying image OCR")
            raw_text, confidence = parse_image(file_path)
            confidence *= 0.7  # Lower confidence for OCR on PDF

    elif ext in (".docx", ".doc"):
        raw_text, confidence = parse_docx(file_path)

    elif ext in (".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tiff"):
        raw_text, confidence = parse_image(file_path)

    else:
        raise ValueError(f"Unsupported format: {ext}. Supported: PDF, DOCX, JPG, PNG")

    if not raw_text or len(raw_text.split()) < 10:
        raise ValueError(f"Could not extract readable text from {path.name}")

    # Detect sections
    sections = detect_sections(raw_text)

    # Extract experience data (EXCLUDING education section)
    exp_data = calculate_experience(sections)

    # Keyword stuffing detection
    stuffing_flag = detect_keyword_stuffing(sections, raw_text)

    # Name extraction
    name = extract_name(raw_text)

    # Candidate type classification
    candidate_type = classify_candidate_type(exp_data, sections)

    # Parse confidence adjustment
    skill_count_estimate = len([
        line for line in raw_text.split('\n')
        if len(line.strip()) > 3
    ])
    final_confidence = min(1.0, confidence + (min(skill_count_estimate, 50) / 100))

    return {
        "raw_text": raw_text,
        "sections": sections,
        "name": name,
        "parse_confidence": round(final_confidence, 2),
        "keyword_stuffing_flag": stuffing_flag,
        "format": ext,
        "year_ranges": exp_data["year_ranges"],
        "total_experience_years": exp_data["total_experience_years"],
        "most_recent_year": exp_data.get("most_recent_year"),
        "has_professional_experience": exp_data["has_professional_experience"],
        "candidate_type": candidate_type,
        "section_keys_found": list(sections.keys())
    }


def build_candidate_profile(file_path: str, graph: dict) -> dict:
    """
    Full pipeline: file → complete candidate profile with skills.
    Called at resume upload time.
    """
    from alias_matcher import match_skills_to_graph

    parsed = parse_resume(file_path)

    # Extract skills per section with context
    skill_nodes = []
    seen_nodes = set()

    DEPTH_SIGNALS = {
        "lead": ["led", "leading", "lead", "headed", "managed", "owned",
                 "architected", "designed", "spearheaded", "directed"],
        "contribute": ["built", "developed", "implemented", "created",
                      "contributed", "wrote", "shipped", "deployed",
                      "maintained", "improved", "worked on"],
        "mention": ["familiar", "exposure", "knowledge of", "experience with",
                   "understanding of", "aware of", "learning", "studied"]
    }

    CONTEXT_MAP = {
        "experience": "professional",
        "professional experience": "professional",
        "projects": "academic",
        "personal projects": "personal",
        "academic projects": "academic",
        "skills": "mention",
        "technical skills": "mention",
        "certifications": "cert",
        "education": "academic",
        "general": "mention"
    }

    for section_name, section_text in parsed["sections"].items():
        if not section_text.strip():
            continue

        matched = match_skills_to_graph(section_text, graph)
        context = CONTEXT_MAP.get(section_name.lower(), "mention")

        # Infer depth from action verbs
        text_lower = section_text.lower()
        depth = "mention"
        for d, signals in DEPTH_SIGNALS.items():
            if any(s in text_lower for s in signals):
                depth = d
                break

        for skill in matched:
            if skill["node_id"] not in seen_nodes:
                skill_nodes.append({
                    **skill,
                    "depth": depth,
                    "context": context,
                    "section": section_name
                })
                seen_nodes.add(skill["node_id"])

    # Skill verification
    skill_verification = verify_skills_against_evidence(
        parsed["sections"], graph
    )

    return {
        "candidate_id": str(uuid.uuid4()),
        "name": parsed["name"],
        "email": None,
        "resume_format": parsed["format"],
        "parse_confidence": parsed["parse_confidence"],
        "keyword_stuffing_flag": parsed["keyword_stuffing_flag"],
        "raw_text": parsed["raw_text"],
        "sections": parsed["sections"],
        "skill_nodes": skill_nodes,
        "skill_node_ids": list(seen_nodes),
        "year_ranges": parsed["year_ranges"],
        "total_experience_years": parsed["total_experience_years"],
        "most_recent_year": parsed["most_recent_year"],
        "has_professional_experience": parsed["has_professional_experience"],
        "candidate_type": parsed["candidate_type"],
        "skill_verification": skill_verification,
        "section_keys_found": parsed["section_keys_found"]
    }
